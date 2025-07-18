import streamlit as st
import gspread
from google.oauth2.service_account import Credentials
from openai import OpenAI
from docx import Document
from docx.shared import Pt
from io import BytesIO
from datetime import date
import pandas as pd
import altair as alt
import datetime

# === AUTH + PAGE CONFIG ===
st.set_page_config(page_title="Mestek Coaching Generator", layout="wide")
PASSWORD = "WFHQmestek413"
if "authenticated" not in st.session_state:
    st.session_state.authenticated = False

if not st.session_state.authenticated:
    with st.form("password_form"):
        pw_input = st.text_input("Enter password", type="password")
        if st.form_submit_button("Unlock") and pw_input == PASSWORD:
            st.session_state.authenticated = True
            st.rerun()
        elif pw_input:
            st.error("Incorrect password.")
    st.stop()

# === GOOGLE SETUP ===
scope = ['https://spreadsheets.google.com/feeds', 'https://www.googleapis.com/auth/drive']
service_account_info = st.secrets["gcp_service_account"]
creds = Credentials.from_service_account_info(service_account_info, scopes=scope)
gs_client = gspread.authorize(creds)
sheet = gs_client.open("Coaching Tracker").sheet1  # <-- rename to match your sheet

# === OPENAI ===
client = OpenAI(api_key=st.secrets["openai"]["api_key"])

# === HELPERS ===
def parse_coaching_sections(raw_text):
    sections = {}
    current = None
    buffer = []
    for line in raw_text.splitlines():
        line = line.strip()
        if line.endswith(":") and line[:-1] in ["Incident Summary", "Expectations Going Forward", "Tags", "Severity"]:
            if current and buffer:
                sections[current] = " ".join(buffer).strip()
                buffer = []
            current = line[:-1]
        elif current:
            buffer.append(line)
    if current and buffer:
        sections[current] = " ".join(buffer).strip()
    return sections

def build_doc(latest, coaching_dict):
    doc = Document()
    doc.add_heading("Employee Coaching & Counseling Form", 0)
    doc.add_paragraph(f"(Created {date.today().strftime('%m/%d/%y')})")
    doc.add_heading("Section 1 – Supervisor Entry", level=1)
    for field in [
        "Date of Incident", "Department", "Employee Name", "Supervisor Name",
        "Action to be Taken", "Issue Type", "Incident Description", "Estimated/Annual Cost",
        "Language Spoken", "Previous Coaching/Warnings"]:
        para = doc.add_paragraph()
        para.add_run(field + ": ").bold = True
        para.add_run(latest.get(field, "[Missing]"))

    doc.add_page_break()
    doc.add_heading("Section 2 – AI-Generated Coaching Report", level=1)
    for section in ["Incident Summary", "Expectations Going Forward", "Tags", "Severity"]:
        if section in coaching_dict:
            doc.add_paragraph(section + ":", style="Heading 2")
            doc.add_paragraph(coaching_dict[section])
    return doc

def log_submission_to_sheet(data_dict):
    timestamp = datetime.datetime.now().strftime("%m/%d/%Y %H:%M:%S")
    row = [
        timestamp,
        data_dict.get("Supervisor Name", ""),
        data_dict.get("Employee Name", ""),
        data_dict.get("Department", ""),
        data_dict.get("Date of Incident", ""),
        data_dict.get("Issue Type", ""),
        data_dict.get("Action to be Taken", ""),
        data_dict.get("Incident Description", ""),
        data_dict.get("Estimated/Annual Cost", ""),
        data_dict.get("Language Spoken", ""),
        data_dict.get("Previous Coaching/Warnings", "")
    ]
    sheet.append_row(row, value_input_option="USER_ENTERED")

# === FORM ===
tab1, tab2 = st.tabs(["📝 Coaching Form", "📊 Trend Dashboard"])

with tab1:
    with st.form("coaching_form"):
        supervisor = st.selectbox("Supervisor Name", [
            "Marty", "Nick", "Pete", "Ralph", "Steve", "Bill", "John",
            "Janitza", "Fundi", "Lisa", "Dave", "Dean"])
        employee = st.text_input("Employee Name")
        department = st.selectbox("Department", [
            "Rough In", "Paint Line (NP)", "Commercial Fabrication",
            "Baseboard Accessories", "Maintenance", "Residential Fabrication",
            "Residential Assembly/Packing", "Warehouse (55WIPR)",
            "Convector & Twin Flo", "Shipping/Receiving/Drivers",
            "Dadanco Fabrication/Assembly", "Paint Line (Dadanco)"])
        incident_date = st.date_input("Date of Incident", value=date.today())
        issue_type = st.selectbox("Issue Type", [
            "Attendance", "Safety", "Behavior", "Performance", "Policy Violation", "Recognition"])
        action_taken = st.selectbox("Action to be Taken", [
            "Coaching", "Verbal Warning", "Written Warning", "Suspension", "Termination"])
        description = st.text_area("Incident Description")
        estimated_cost = st.text_input("Estimated/Annual Cost (optional)")
        lang_opt = st.selectbox("Language Spoken", ["English", "Spanish", "Other"])
        language = st.text_input("Please specify the language:") if lang_opt == "Other" else lang_opt
        previous = st.text_area("Previous Coaching/Warnings (if any)")
        submitted = st.form_submit_button("Generate Coaching Report")

    if submitted:
        latest = {
            "Timestamp": date.today().isoformat(),
            "Supervisor Name": supervisor,
            "Employee Name": employee,
            "Department": department,
            "Date of Incident": incident_date.strftime("%Y-%m-%d"),
            "Issue Type": issue_type,
            "Action to be Taken": action_taken,
            "Incident Description": description,
            "Estimated/Annual Cost": estimated_cost,
            "Language Spoken": language,
            "Previous Coaching/Warnings": previous
        }

        prompt = f"""
You are a workplace coaching assistant. Generate a coaching report in sections:
Incident Summary:
Expectations Going Forward:
Tags:
Severity:

Data:
Supervisor: {supervisor}
Employee: {employee}
Department: {department}
Date: {incident_date}
Issue: {issue_type}
Action: {action_taken}
Description: {description}
"""

        with st.spinner("Generating coaching report..."):
            response = client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[{"role": "user", "content": prompt}]
            ).choices[0].message.content.strip()

            if language.lower() != "english":
                response = client.chat.completions.create(
                    model="gpt-3.5-turbo",
                    messages=[{"role": "user", "content": f"Translate into {language}:\n{response}"}]
                ).choices[0].message.content.strip()

        coaching_sections = parse_coaching_sections(response)
        doc = build_doc(latest, coaching_sections)
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        try:
            log_submission_to_sheet(latest)
            st.success("✅ Logged to Google Sheet")
        except Exception as e:
            st.error(f"❌ Google Sheet error: {e}")

        st.download_button("📄 Download Coaching Report", data=buffer,
                           file_name=f"{employee.replace(' ', '_')}_coaching.docx")

# === TREND DASHBOARD ===
with tab2:
    st.header("📊 Coaching Trend Dashboard")
    try:
        df = pd.DataFrame(sheet.get_all_records())
        df["Date of Incident"] = pd.to_datetime(df["Date of Incident"], errors="coerce")

        min_date = df["Date of Incident"].min()
        max_date = df["Date of Incident"].max()
        start_date, end_date = st.date_input("Filter by Date Range", [min_date, max_date], key="date_range_filter")

        if start_date and end_date:
            df = df[(df["Date of Incident"] >= pd.to_datetime(start_date)) &
                    (df["Date of Incident"] <= pd.to_datetime(end_date))]

        filter_action = st.selectbox("Filter by Action Taken",
                                     ["All"] + df["Action to be Taken"].dropna().unique().tolist())
        if filter_action != "All":
            df = df[df["Action to be Taken"] == filter_action]

        st.dataframe(df)

        st.subheader("Issue Type Count")
        issue_counts = df["Issue Type"].value_counts().reset_index()
        issue_counts.columns = ["Issue Type", "Count"]
        chart = alt.Chart(issue_counts).mark_bar().encode(
            x=alt.X("Issue Type", sort="-y"),
            y="Count",
            tooltip=["Issue Type", "Count"]
        ).properties(width=600, height=400)
        st.altair_chart(chart, use_container_width=True)

        st.subheader("Actions Over Time")
        df["Date Only"] = df["Date of Incident"].dt.date
        trend = df.groupby(["Date Only", "Action to be Taken"]).size().unstack(fill_value=0)
        st.line_chart(trend)

    except Exception as e:
        st.error(f"❌ Could not load dashboard: {e}")
